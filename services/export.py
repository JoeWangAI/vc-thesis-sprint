"""
Export service - Generate Word memos and enhanced CSV exports.
"""
from typing import List, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import Company, ThesisSprint, ShortlistEntry


class ExportService:
    """Service for generating export artifacts."""

    def generate_word_memo(
        self,
        sprint: ThesisSprint,
        shortlist: List[Tuple[Company, ShortlistEntry]]
    ) -> Document:
        """
        Generate Word memo document.

        Structure:
        1. Executive Summary Table (10 companies)
        2. Thesis Framing
        3. Shortlist Detail (10 mini-cards per company)
        4. Appendix (sources + conflict resolution notes)

        Args:
            sprint: The thesis sprint
            shortlist: List of (company, shortlist_entry) tuples

        Returns:
            python-docx Document object
        """
        doc = Document()

        # Set up document styles
        self._setup_document_styles(doc)

        # Title
        title = doc.add_heading(f'Investment Memo: {sprint.name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Section 1: Executive Summary Table
        self._add_executive_summary(doc, shortlist[:10])

        doc.add_page_break()

        # Section 2: Thesis Framing
        self._add_thesis_framing(doc, sprint)

        doc.add_page_break()

        # Section 3: Shortlist Detail
        self._add_shortlist_detail(doc, shortlist[:10])

        doc.add_page_break()

        # Section 4: Appendix
        self._add_appendix(doc, shortlist[:10])

        return doc

    def _setup_document_styles(self, doc: Document):
        """Set up document-level styles."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_executive_summary(
        self,
        doc: Document,
        shortlist: List[Tuple[Company, ShortlistEntry]]
    ):
        """Add executive summary table."""
        doc.add_heading('Executive Summary', 1)

        if not shortlist:
            doc.add_paragraph('No companies in shortlist.')
            return

        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['Company', 'Stage', 'Last Round', 'Amount', 'Confidence', 'Status']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            self._make_bold(hdr_cells[i].paragraphs[0])

        # Data rows
        for company, entry in shortlist:
            row_cells = table.add_row().cells
            row_cells[0].text = company.name

            # Stage
            if company.stage_estimate:
                row_cells[1].text = company.stage_estimate.stage
            else:
                row_cells[1].text = company.stage or 'Unknown'

            # Last round info
            if company.funding_snapshot:
                row_cells[2].text = company.funding_snapshot.last_round_type or 'N/A'
                row_cells[3].text = company.funding_snapshot.amount or 'N/A'
                conf = company.funding_snapshot.overall_confidence.value.capitalize()
                row_cells[4].text = conf
            elif company.funding_events:
                latest = company.funding_events[0]
                row_cells[2].text = latest.round_type
                row_cells[3].text = latest.amount or 'N/A'
                row_cells[4].text = company.confidence.value.capitalize()
            else:
                row_cells[2].text = 'N/A'
                row_cells[3].text = 'N/A'
                row_cells[4].text = 'N/A'

            row_cells[5].text = entry.status.value.capitalize()

    def _add_thesis_framing(self, doc: Document, sprint: ThesisSprint):
        """Add thesis framing section."""
        doc.add_heading('Thesis', 1)

        # Thesis description
        doc.add_heading('Overview', 2)
        doc.add_paragraph(sprint.description)

        # Criteria
        doc.add_heading('Investment Criteria', 2)
        criteria_list = doc.add_paragraph(style='List Bullet')
        criteria_list.add_run(f'Stage Focus: {sprint.stage_focus}')

        criteria_list = doc.add_paragraph(style='List Bullet')
        criteria_list.add_run(f'Geography: {sprint.geography}')

        criteria_list = doc.add_paragraph(style='List Bullet')
        criteria_list.add_run(f'Last Raise Filter: {sprint.last_raise_filter}')

        if sprint.keywords_include:
            criteria_list = doc.add_paragraph(style='List Bullet')
            criteria_list.add_run(f'Include Keywords: {", ".join(sprint.keywords_include)}')

    def _add_shortlist_detail(
        self,
        doc: Document,
        shortlist: List[Tuple[Company, ShortlistEntry]]
    ):
        """Add shortlist detail section with per-company mini-cards."""
        doc.add_heading('Shortlisted Companies', 1)

        for i, (company, entry) in enumerate(shortlist, 1):
            # Company name as heading
            doc.add_heading(f'{i}. {company.name}', 2)

            # What they do
            doc.add_heading('Overview', 3)
            doc.add_paragraph(company.description)

            # Why it fits
            doc.add_heading('Thesis Fit', 3)
            if company.fit_reasons:
                for reason in company.fit_reasons:
                    p = doc.add_paragraph(reason, style='List Bullet')
            elif company.thesis_fit_notes:
                doc.add_paragraph(company.thesis_fit_notes)
            else:
                doc.add_paragraph('No thesis fit notes available.')

            # Funding snapshot
            doc.add_heading('Funding Context', 3)
            if company.funding_snapshot:
                fs = company.funding_snapshot
                para = doc.add_paragraph()
                para.add_run('Last Round: ').bold = True
                para.add_run(f'{fs.last_round_type or "N/A"}\n')

                para.add_run('Date: ').bold = True
                if fs.last_round_date:
                    para.add_run(f'{fs.last_round_date.strftime("%B %Y")}\n')
                else:
                    para.add_run('N/A\n')

                para.add_run('Amount: ').bold = True
                para.add_run(f'{fs.amount or "N/A"}\n')

                para.add_run('Lead: ').bold = True
                para.add_run(f'{fs.lead_investor or "N/A"}\n')

                if fs.post_money_valuation:
                    para.add_run('Valuation: ').bold = True
                    para.add_run(f'{fs.post_money_valuation} ({fs.valuation_basis})\n')

                para.add_run('Confidence: ').bold = True
                para.add_run(f'{fs.overall_confidence.value.capitalize()}')

            elif company.funding_events:
                latest = company.funding_events[0]
                para = doc.add_paragraph()
                para.add_run('Last Round: ').bold = True
                para.add_run(f'{latest.round_type}\n')
                para.add_run('Date: ').bold = True
                para.add_run(f'{latest.date.strftime("%B %Y")}\n')
                para.add_run('Amount: ').bold = True
                para.add_run(f'{latest.amount or "Undisclosed"}')
            else:
                doc.add_paragraph('No funding information available.')

            # Open questions
            doc.add_heading('Open Questions', 3)
            if company.next_action:
                doc.add_paragraph(company.next_action, style='List Bullet')
            else:
                doc.add_paragraph('TBD - requires deeper diligence', style='List Bullet')

            # Separator
            if i < len(shortlist):
                doc.add_paragraph('_' * 80)

    def _add_appendix(
        self,
        doc: Document,
        shortlist: List[Tuple[Company, ShortlistEntry]]
    ):
        """Add appendix with sources and conflict notes."""
        doc.add_heading('Appendix: Sources & Validation Notes', 1)

        for company, entry in shortlist:
            doc.add_heading(company.name, 2)

            # Sources
            if company.funding_snapshot and company.funding_snapshot.sources:
                doc.add_heading('Sources', 3)
                for source in company.funding_snapshot.sources:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{source.title}\n')
                    p.add_run(f'{source.url}\n').font.size = Pt(9)
                    p.add_run(f'Type: {source.source_type}').font.size = Pt(9)

            # Resolution notes
            if company.funding_snapshot and company.funding_snapshot.resolution_note:
                doc.add_heading('Conflict Resolution', 3)
                doc.add_paragraph(company.funding_snapshot.resolution_note)

    def _make_bold(self, paragraph):
        """Make paragraph text bold."""
        for run in paragraph.runs:
            run.bold = True


# Global instance
export_service = ExportService()
